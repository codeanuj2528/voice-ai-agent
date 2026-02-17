import os
import io
import logging
import httpx

import firebase_admin
from firebase_admin import credentials, firestore
from google.cloud.firestore_v1.vector import Vector
from google.cloud.firestore_v1.base_vector_query import DistanceMeasure

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

from config import settings

logger = logging.getLogger(__name__)

# Firestore client (initialized lazily)
_db = None

# Collection name for document chunks
CHUNKS_COLLECTION = "kb_chunks"
DOCUMENTS_COLLECTION = "kb_documents"

# Jina AI embedding config
JINA_EMBEDDING_MODEL = "jina-embeddings-v3"
JINA_EMBEDDING_DIM = 1024
JINA_API_URL = "https://api.jina.ai/v1/embeddings"

# LangChain text splitter
text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=500,
    chunk_overlap=50,
    length_function=len,
    separators=["\n\n", "\n", ". ", " ", ""],
)


def _get_db():
    """Get or initialize Firestore client using Firebase Admin SDK."""
    global _db
    if _db is None:
        if not firebase_admin._apps:
            cred = credentials.Certificate({
                "type": "service_account",
                "project_id": settings.firebase_project_id,
                "client_email": settings.firebase_client_email,
                "private_key": settings.firebase_private_key.replace("\\n", "\n"),
                "token_uri": "https://oauth2.googleapis.com/token",
            })
            firebase_admin.initialize_app(cred)
        _db = firestore.client()
    return _db


async def _get_embeddings(texts: list[str]) -> list[list[float]]:
    """Generate embeddings using Jina AI API."""
    async with httpx.AsyncClient() as client:
        response = await client.post(
            JINA_API_URL,
            headers={
                "Authorization": f"Bearer {settings.jina_api_key}",
                "Content-Type": "application/json",
            },
            json={
                "model": JINA_EMBEDDING_MODEL,
                "input": texts,
                "task": "retrieval.passage",
                "dimensions": JINA_EMBEDDING_DIM,
            },
            timeout=60.0,
        )
        response.raise_for_status()
        data = response.json()
        return [item["embedding"] for item in data["data"]]


async def _get_query_embedding(query: str) -> list[float]:
    """Generate embedding for a search query."""
    async with httpx.AsyncClient() as client:
        response = await client.post(
            JINA_API_URL,
            headers={
                "Authorization": f"Bearer {settings.jina_api_key}",
                "Content-Type": "application/json",
            },
            json={
                "model": JINA_EMBEDDING_MODEL,
                "input": [query],
                "task": "retrieval.query",
                "dimensions": JINA_EMBEDDING_DIM,
            },
            timeout=30.0,
        )
        response.raise_for_status()
        data = response.json()
        return data["data"][0]["embedding"]


def _load_documents(content: bytes, ext: str, filename: str) -> list[Document]:
    """
    Load file bytes into LangChain Document objects with rich metadata.
    Each page (for PDFs) or the full content becomes a Document with metadata.
    """
    if ext in (".txt", ".md"):
        text = content.decode("utf-8")
        return [Document(
            page_content=text,
            metadata={
                "source": filename,
                "file_type": ext.lstrip("."),
                "page": 1,
                "total_pages": 1,
            },
        )]

    elif ext == ".pdf":
        from PyPDF2 import PdfReader
        reader = PdfReader(io.BytesIO(content))
        total_pages = len(reader.pages)
        docs = []
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text and page_text.strip():
                docs.append(Document(
                    page_content=page_text,
                    metadata={
                        "source": filename,
                        "file_type": "pdf",
                        "page": i + 1,
                        "total_pages": total_pages,
                    },
                ))
        return docs

    elif ext == ".docx":
        from docx import Document as DocxDocument
        doc = DocxDocument(io.BytesIO(content))
        text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        return [Document(
            page_content=text,
            metadata={
                "source": filename,
                "file_type": "docx",
                "page": 1,
                "total_pages": 1,
            },
        )]

    else:
        raise ValueError(f"Unsupported file type: {ext}")


async def ingest_document(content: bytes, ext: str, doc_id: str, filename: str):
    """
    Ingest a document into the knowledge base (in-memory).

    Steps:
    1. Load bytes into LangChain Document objects (with metadata)
    2. Split documents using RecursiveCharacterTextSplitter (metadata preserved)
    3. Generate embeddings via Jina AI
    4. Store chunks + embeddings + metadata in Firestore
    """
    logger.info(f"Starting ingestion for '{filename}' (doc_id={doc_id})")

    # 1. Load into LangChain Documents
    try:
        docs = _load_documents(content, ext, filename)
        if not docs:
            raise ValueError("Document contains no extractable text")
        total_chars = sum(len(d.page_content) for d in docs)
        logger.info(f"[{filename}] Loaded {len(docs)} page(s), {total_chars} characters")
    except Exception:
        logger.exception(f"[{filename}] Document loading failed")
        raise

    # 2. Split documents (metadata is carried over to each chunk)
    try:
        split_docs = text_splitter.split_documents(docs)
        if not split_docs:
            raise ValueError("No chunks generated from document")
        logger.info(f"[{filename}] Split into {len(split_docs)} chunks (chunk_size=500, overlap=50)")
    except Exception:
        logger.exception(f"[{filename}] Text splitting failed")
        raise

    # 3. Generate embeddings (batch)
    try:
        batch_size = 50
        all_embeddings = []
        for i in range(0, len(split_docs), batch_size):
            batch_texts = [doc.page_content for doc in split_docs[i:i + batch_size]]
            embeddings = await _get_embeddings(batch_texts)
            all_embeddings.extend(embeddings)
            logger.info(f"[{filename}] Embedded batch {i // batch_size + 1} ({len(batch_texts)} chunks)")

        logger.info(f"[{filename}] Generated {len(all_embeddings)} embeddings total")
    except Exception:
        logger.exception(f"[{filename}] Embedding generation failed")
        raise

    # 4. Store in Firestore
    try:
        db = _get_db()

        # Store document metadata
        doc_ref = db.collection(DOCUMENTS_COLLECTION).document(doc_id)
        doc_ref.set({
            "doc_id": doc_id,
            "filename": filename,
            "file_type": ext.lstrip("."),
            "chunk_count": len(split_docs),
            "total_pages": docs[0].metadata.get("total_pages", 1),
            "status": "ingested",
        })
        logger.info(f"[{filename}] Stored document metadata in '{DOCUMENTS_COLLECTION}'")

        # Store chunks with embeddings and metadata
        batch_writer = db.batch()
        for i, (doc, embedding) in enumerate(zip(split_docs, all_embeddings)):
            chunk_ref = db.collection(CHUNKS_COLLECTION).document(f"{doc_id}_chunk_{i}")
            batch_writer.set(chunk_ref, {
                "doc_id": doc_id,
                "chunk_index": i,
                "content": doc.page_content,
                "metadata": {
                    "source": doc.metadata.get("source", filename),
                    "file_type": doc.metadata.get("file_type", ""),
                    "page": doc.metadata.get("page", 1),
                    "total_pages": doc.metadata.get("total_pages", 1),
                    "loc": doc.metadata.get("loc", {}),
                },
                "embedding": Vector(embedding),
            })

            # Firestore batch limit is 500
            if (i + 1) % 400 == 0:
                batch_writer.commit()
                batch_writer = db.batch()
                logger.info(f"[{filename}] Committed batch of 400 chunks to Firestore")

        batch_writer.commit()
        logger.info(f"[{filename}] ✅ Ingestion complete — {len(split_docs)} chunks stored in '{CHUNKS_COLLECTION}'")

    except Exception:
        logger.exception(f"[{filename}] Firestore storage failed")
        raise


async def delete_document_chunks(doc_id: str):
    """Delete all chunks for a document from Firestore."""
    db = _get_db()

    # Delete chunks
    chunks_ref = db.collection(CHUNKS_COLLECTION)
    query = chunks_ref.where("doc_id", "==", doc_id)
    docs = query.stream()

    batch = db.batch()
    count = 0
    for doc in docs:
        batch.delete(doc.reference)
        count += 1
        if count % 400 == 0:
            batch.commit()
            batch = db.batch()

    batch.commit()

    # Delete document metadata
    db.collection(DOCUMENTS_COLLECTION).document(doc_id).delete()

    logger.info(f"Deleted {count} chunks for document {doc_id}")


async def list_documents() -> list[dict]:
    """List all documents in the knowledge base."""
    db = _get_db()
    docs_ref = db.collection(DOCUMENTS_COLLECTION)
    docs = docs_ref.stream()

    result = []
    for doc in docs:
        data = doc.to_dict()
        result.append({
            "doc_id": data.get("doc_id", doc.id),
            "filename": data.get("filename", "unknown"),
            "status": data.get("status", "unknown"),
        })

    return result
